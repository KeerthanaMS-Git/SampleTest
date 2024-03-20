import nltk
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from collections import Counter
import string
import easyocr
import easyocr
from pdf2image import convert_from_path
from PIL import Image, ImageFilter
from docx import Document
import os
import numpy as np

def preprocess_image(image):
    # Convert the image to grayscale
    grayscale_image = image.convert('L')
    
    # Apply any additional preprocessing steps here, such as noise reduction, thresholding, etc.
    
    # Return the preprocessed image
    return grayscale_image

def extract_text_from_file(input_path):
    # Initialize the EasyOCR reader
    reader = easyocr.Reader(['en'])  # You can specify other languages as needed

    # Initialize a new Word document
    doc = Document()

    # Check if the input file is a PDF
    if input_path.lower().endswith('.pdf'):
        # Convert PDF to images
        images = convert_from_path(input_path)

        # Extract text from each page
        for i, image in enumerate(images):
            print(f'Extracting text from page {i + 1}...')
            
            # Preprocess the image
            preprocessed_image = preprocess_image(image)
            image_np = np.array(preprocessed_image)
            
            # Perform OCR on the preprocessed image
            result = reader.readtext(image_np)
            print("OCR Result:", result)
            output_text = ""
            unanalyzed_words = 0
            for idx, word in enumerate(result):
                if idx > 0:
                    prev_word = result[idx - 1]
                    if word[2] < 0.2:
                        unanalyzed_words += 1
                    if word[0][3][1] - prev_word[0][3][1] < 15:
                        n_spaces = (word[0][0][0] - prev_word[0][1][0])/20
                        n_spaces = int(n_spaces)
                        spaces = n_spaces * ' '
                        output_text += spaces + word[1]
                        # output_text += ' ' + word[1]
                    else:
                        output_text += "\n" + word[1]
                else:
                    output_text += word[1]
            # Extracted text
            print("Output text is :", output_text)
            return output_text
        
def document_analysis(input_path,output_path):
# Sample text extracted from the document (replace with your extracted text)
    document_text = extract_text_from_file(input_path)

    # Tokenization
    tokens = word_tokenize(document_text.lower())

    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    tokens = [word for word in tokens if word not in stop_words]

    # Remove punctuation
    tokens = [word for word in tokens if word not in string.punctuation]

    # Stemming
    stemmer = PorterStemmer()
    tokens = [stemmer.stem(word) for word in tokens]

    # Count word frequencies
    word_freq = Counter(tokens)

    # Print most common words
    print("Most common words:")
    print(word_freq.most_common(10))

    # Sentence Tokenization
    sentences = sent_tokenize(document_text)

    # Print number of sentences
    print("\nNumber of sentences:", len(sentences))

    # Perform Named Entity Recognition (NER) using spaCy or NLTK's chunking techniques
    # Example using NLTK chunking:
    grammar = r'NP: {<DT>?<JJ>*<NN>}'
    chunk_parser = nltk.RegexpParser(grammar)
    tagged_tokens = nltk.pos_tag(tokens)
    chunks = chunk_parser.parse(tagged_tokens)


    doc = Document()
    
    # Write most common words
    doc.add_heading('Most Common Words', level=1)
    for word, freq in word_freq.most_common(10):
        doc.add_paragraph(f'{word}: {freq}')

    # Write number of sentences
    doc.add_heading('Number of Sentences', level=1)
    doc.add_paragraph(str(len(sentences)))

    # Write named entities
    doc.add_heading('Named Entities', level=1)
    named_entities = []
    for subtree in chunks.subtrees():
        if subtree.label() == 'NP':
            named_entities.append(' '.join(word for word, pos in subtree.leaves()))
    for entity in named_entities:
        doc.add_paragraph(entity)

    # Save the document
    doc.save(output_path)
    # Print named entities
    print("\nNamed Entities:")
    for subtree in chunks.subtrees():
        if subtree.label() == 'NP':
            print(' '.join(word for word, pos in subtree.leaves()))


input_path = r'SampleInputs\PDF\Demo - 1-1.pdf' 
output_path = r'OutputFiles\EasyOCROutput\DocumentAnalysis.doc' 
document_analysis(input_path,output_path)

# Perform more advanced analysis as needed (e.g., sentiment analysis, document classification)
