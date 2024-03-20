import easyocr
from pdf2image import convert_from_path
from PIL import Image, ImageFilter
from docx import Document
import os
import numpy as np

def preprocess_image(image):
    # Convert the image to grayscale
    grayscale_image = image.convert('L')
    
    # Apply any additional preprocessing steps here, such as noise reduction, thresholding, etc.
    
    # Return the preprocessed image
    return grayscale_image

def extract_text_from_file(input_path, output_path):
    # Initialize the EasyOCR reader
    reader = easyocr.Reader(['en'])  # You can specify other languages as needed

    # Initialize a new Word document
    doc = Document()

    # Check if the input file is a PDF
    if input_path.lower().endswith('.pdf'):
        # Convert PDF to images
        images = convert_from_path(input_path)

        # Extract text from each page
        for i, image in enumerate(images):
            print(f'Extracting text from page {i + 1}...')
            
            # Preprocess the image
            preprocessed_image = preprocess_image(image)
            image_np = np.array(preprocessed_image)
            
            # Perform OCR on the preprocessed image
            result = reader.readtext(image_np)
            print("OCR Result:", result)
            output_text = ""
            unanalyzed_words = 0
            for idx, word in enumerate(result):
                if idx > 0:
                    prev_word = result[idx - 1]
                    if word[2] < 0.2:
                        unanalyzed_words += 1
                    if word[0][3][1] - prev_word[0][3][1] < 15:
                        n_spaces = (word[0][0][0] - prev_word[0][1][0])/20
                        n_spaces = int(n_spaces)
                        spaces = n_spaces * ' '
                        output_text += spaces + word[1]
                        # output_text += ' ' + word[1]
                    else:
                        output_text += "\n" + word[1]
                else:
                    output_text += word[1]
            # Extracted text
            print("Output text is :", output_text)
            print("\n"+"Total unanalyzed words = "+str(unanalyzed_words))
            doc.add_paragraph(output_text)

            # Add a newline between each page
            if i < len(images) - 1:
                doc.add_paragraph('\n')
            # Rest of the code remains the same as before...
            # (Omitted for brevity)


    # Check if the input file is an image
    elif any(input_path.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']):
        # Read text from the image
        print('Extracting text from image...')
        
        # Open the image
        image = Image.open(input_path)
        
        # Preprocess the image
        preprocessed_image = preprocess_image(image)
        image_np = np.array(preprocessed_image)
        
        # Perform OCR on the preprocessed image
        result = reader.readtext(image_np)
        print("OCR Result:", result)

        # Rest of the code remains the same as before...
        # (Omitted for brevity)

    else:
        print('Unsupported file format.')
        return

    # Save the Word document to the custom output path
    output_docx = output_path
    doc.save(output_docx)
    print(f'Extracted text has been saved to {output_docx}')

# Path to your input file (PDF or image)
input_path = r'SampleInputs\PDF\Demo - 1-1.pdf'   # Change this to your file path

# Custom output path
output_path = r'OutputFiles\EasyOCROutput\extracted_text.doc'

# Extract text from the input file and save as Word document to custom output path
extract_text_from_file(input_path, output_path)
